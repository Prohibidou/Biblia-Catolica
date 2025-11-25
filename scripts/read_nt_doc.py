import sys

try:
    # Intentar con python-docx primero
    from docx import Document
    
    doc_path = r"BibliaPDF/NT Eunsa con notas (1).doc"
    
    print("Intentando abrir con python-docx...")
    doc = Document(doc_path)
    
    print(f"Documento abierto exitosamente!")
    print(f"Número de párrafos: {len(doc.paragraphs)}")
    
    # Mostrar primeros 20 párrafos
    print("\nPrimeros 20 párrafos:")
    print("="*70)
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"{i+1}. {para.text[:100]}")
    
    # Guardar muestra
    with open('scripts/nt_doc_muestra.txt', 'w', encoding='utf-8') as f:
        for i, para in enumerate(doc.paragraphs[:200]):
            if para.text.strip():
                f.write(f"{i+1}. {para.text}\n")
    
    print("\nMuestra guardada en scripts/nt_doc_muestra.txt")
    
except ImportError:
    print("python-docx no está instalado. Instalando...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    print("\n¡Instalado! Ejecuta el script nuevamente.")
except Exception as e:
    print(f"Error al abrir como .docx: {e}")
    print("\nEl archivo puede ser .doc (formato antiguo), no .docx")
    print("Intentando con método alternativo...")
    
    # Intentar leer como texto plano o con otro método
    try:
        import zipfile
        with zipfile.ZipFile(doc_path, 'r') as zip_ref:
            print("Es un archivo comprimido")
    except:
        print("No es formato .docx (Office Open XML)")
        print("\nNecesitamos convertir .doc a .docx o usar otra herramienta.")
        print("\nOpciones:")
        print("1. Abrir el archivo en Word y guardarlo como .docx")
        print("2. Usar pywin32 para leer .doc (solo Windows)")
        print("3. Convertir en línea o con LibreOffice")
